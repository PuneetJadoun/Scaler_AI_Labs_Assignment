"""Reads .docx files and converts them into structured Document objects.

This module is the single point of contact with python-docx in the entire
application. Every other module (detectors, replacement, evaluation, and
eventually the writer) works exclusively with the plain dataclasses defined
in ``document.models`` and must never import ``docx`` directly.
"""

from pathlib import Path
from typing import Any

import docx
from docx.document import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from document.models.cell import Cell, Row
from document.models.document import Document
from document.models.paragraph import Paragraph
from document.models.table import Table

_DOCX_SUFFIX = ".docx"


class DocumentReaderError(Exception):
    """Base exception for all failures raised while reading a document."""


class InvalidDocumentPathError(DocumentReaderError):
    """Raised when the given path does not point to a readable file."""


class UnsupportedDocumentFormatError(DocumentReaderError):
    """Raised when the file does not have a .docx extension."""


class CorruptedDocumentError(DocumentReaderError):
    """Raised when the file cannot be opened as a valid Word document."""


class EmptyDocumentError(DocumentReaderError):
    """Raised when the document contains no paragraphs and no tables."""


def read_document(path: str | Path) -> Document:
    """Read a .docx file and return it as a structured Document object.

    Args:
        path: Path to a .docx file, as a string or Path.

    Returns:
        A Document containing every paragraph and table found in the file.

    Raises:
        InvalidDocumentPathError: The path does not exist or is not a file.
        UnsupportedDocumentFormatError: The file is not a .docx file.
        CorruptedDocumentError: The file cannot be parsed as a Word document.
        EmptyDocumentError: The document has no paragraphs and no tables.
    """
    file_path = _validate_path(path)
    docx_file = _open_docx(file_path)

    paragraphs = _read_paragraphs(docx_file)
    tables = _read_tables(docx_file)

    if not paragraphs and not tables:
        raise EmptyDocumentError(f"Document '{file_path}' contains no readable content.")

    return Document(source_path=file_path, paragraphs=paragraphs, tables=tables)


def _validate_path(path: str | Path) -> Path:
    """Resolve and validate that ``path`` points to a readable .docx file."""
    file_path = Path(path)

    if not file_path.exists():
        raise InvalidDocumentPathError(f"File not found: '{file_path}'")

    if not file_path.is_file():
        raise InvalidDocumentPathError(f"Path is not a file: '{file_path}'")

    if file_path.suffix.lower() != _DOCX_SUFFIX:
        raise UnsupportedDocumentFormatError(
            f"Unsupported file format '{file_path.suffix}'. "
            f"Only '{_DOCX_SUFFIX}' files are supported."
        )

    return file_path


def _open_docx(file_path: Path) -> DocxDocument:
    """Open ``file_path`` with python-docx, translating parse failures."""
    try:
        return docx.Document(str(file_path))
    except PackageNotFoundError as exc:
        raise CorruptedDocumentError(
            f"File '{file_path}' is not a valid or is a corrupted .docx document."
        ) from exc


def _read_paragraphs(docx_file: DocxDocument) -> list[Paragraph]:
    """Extract every top-level paragraph, in document order."""
    return [
        Paragraph(id=index, text=paragraph.text)
        for index, paragraph in enumerate(docx_file.paragraphs)
    ]


def _read_tables(docx_file: DocxDocument) -> list[Table]:
    """Extract every table, in document order."""
    return [_read_table(index, table) for index, table in enumerate(docx_file.tables)]


def _read_table(table_id: int, docx_table: Any) -> Table:
    """Convert a single python-docx table into a Table."""
    rows = [_read_row(row_id, row) for row_id, row in enumerate(docx_table.rows)]
    return Table(id=table_id, rows=rows)


def _read_row(row_id: int, docx_row: Any) -> Row:
    """Convert a single python-docx row into a Row."""
    cells = [Cell(id=cell_id, text=cell.text) for cell_id, cell in enumerate(docx_row.cells)]
    return Row(id=row_id, cells=cells)
