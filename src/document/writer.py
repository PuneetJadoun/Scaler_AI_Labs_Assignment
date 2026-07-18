"""Writes a structured Document back out to a .docx file.

This is the only module besides ``document.reader`` permitted to import
python-docx. Its single responsibility is serialization: it takes a
Document that has already been through the Detector and Replacement
modules and writes down exactly the text it finds there. It performs no
PII detection, no replacement, and no fake-data generation of its own -
by the time a Document reaches this module, every decision about what the
text should say has already been made elsewhere.
"""

import logging
import re
from pathlib import Path

import docx

from document.models.document import Document
from document.models.table import Table

logger = logging.getLogger(__name__)

_DOCX_SUFFIX = ".docx"
_REDACTED_SUFFIX = "_redacted"
_DEFAULT_OUTPUT_DIR = Path(__file__).resolve().parent.parent.parent / "output"
_WHITESPACE_RUN = re.compile(r"\s+")


class DocumentWriterError(Exception):
    """Base exception for all failures raised while writing a document."""


class EmptyDocumentError(DocumentWriterError):
    """Raised when the document contains no paragraphs and no tables."""


class InvalidOutputPathError(DocumentWriterError):
    """Raised when the output directory is invalid or cannot be created."""


class DocumentWriteError(DocumentWriterError):
    """Raised when python-docx fails to save the file to disk."""


def write_document(document: Document, output_dir: str | Path | None = None) -> Path:
    """Serialize `document` to a new .docx file and save it to `output_dir`.

    Args:
        document: The structured Document to write, already processed by
            the Reader, Detector, and Replacement modules.
        output_dir: Directory to write the file into. Created automatically
            if it doesn't exist. Defaults to the project's top-level
            ``output/`` directory.

    Returns:
        The path to the written .docx file.

    Raises:
        EmptyDocumentError: `document` has no paragraphs and no tables.
        InvalidOutputPathError: The output path exists but isn't a
            directory, or the directory can't be created (e.g. a
            permission error).
        DocumentWriteError: python-docx failed to save the file (e.g. a
            permission error or disk failure during the write itself).
    """
    if not document.paragraphs and not document.tables:
        raise EmptyDocumentError(
            f"Document from '{document.source_path}' has no paragraphs and no "
            f"tables; nothing to write."
        )

    resolved_output_dir = _resolve_output_dir(output_dir)
    output_path = resolved_output_dir / _redacted_filename(document.source_path)

    logger.info(
        "Writing %d paragraph(s) and %d table(s) to '%s'.",
        len(document.paragraphs),
        len(document.tables),
        output_path,
    )

    docx_file = _build_docx(document)
    _save(docx_file, output_path)

    logger.info("Document written successfully to '%s'.", output_path)
    return output_path


def _resolve_output_dir(output_dir: str | Path | None) -> Path:
    """Resolve and create `output_dir` (or the default output/ directory)."""
    directory = Path(output_dir) if output_dir is not None else _DEFAULT_OUTPUT_DIR

    if directory.exists() and not directory.is_dir():
        raise InvalidOutputPathError(f"Output path '{directory}' exists and is not a directory.")

    try:
        directory.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        raise InvalidOutputPathError(
            f"Could not create output directory '{directory}': {exc}"
        ) from exc

    return directory


def _redacted_filename(source_path: Path) -> str:
    """Derive the output filename, e.g. 'Report.docx' -> 'Report_redacted.docx'."""
    sanitized_stem = _WHITESPACE_RUN.sub("_", source_path.stem.strip())
    return f"{sanitized_stem}{_REDACTED_SUFFIX}{_DOCX_SUFFIX}"


def _build_docx(document: Document) -> "docx.document.Document":
    """Build an in-memory python-docx document mirroring `document`."""
    docx_file = docx.Document()

    for paragraph in document.paragraphs:
        docx_file.add_paragraph(paragraph.text)

    for table in document.tables:
        _write_table(docx_file, table)

    return docx_file


def _write_table(docx_file: "docx.document.Document", table: Table) -> None:
    """Recreate one table, preserving row and cell order."""
    if not table.rows:
        return

    column_count = max(len(row.cells) for row in table.rows)
    if column_count == 0:
        return

    docx_table = docx_file.add_table(rows=len(table.rows), cols=column_count)
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            docx_table.cell(row_index, cell_index).text = cell.text


def _save(docx_file: "docx.document.Document", output_path: Path) -> None:
    """Save `docx_file` to `output_path`, translating write failures."""
    try:
        docx_file.save(str(output_path))
    except OSError as exc:
        raise DocumentWriteError(f"Failed to write '{output_path}': {exc}") from exc
